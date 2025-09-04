from io import BytesIO
import re
import logging
import time
from docx import Document
from docx.shared import Pt, Inches, Twips

# Set up logging
logger = logging.getLogger("contract_generator.docx_converter")

def convert_to_docx(text):
    """
    Convert contract text to a DOCX file
    
    Args:
        text (str): The contract text
        
    Returns:
        bytes: The DOCX file as bytes
    """
    start_time = time.time()
    logger.info("Starting DOCX conversion")
    logger.debug(f"Converting text of length {len(text)} characters to DOCX")
    
    # Create a new Document
    doc = Document()
    
    try:
        # Configure document styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        logger.debug("Document styles and margins configured")
        
        # Split the text into lines
        lines = text.split('\n')
        logger.debug(f"Text split into {len(lines)} lines")
        
        # Track if we're in a numbered list
        in_numbered_list = False
        list_level = 0
        
        # Track document structure for logging
        headings_count = 0
        paragraphs_count = 0
        list_items_count = 0
        
        for i, line in enumerate(lines):
            line = line.rstrip()
            
            # Skip empty lines
            if not line:
                doc.add_paragraph()
                continue
            
            # Check if the line is a heading (all caps or ends with a colon)
            if line.isupper() or (line.endswith(':') and len(line) < 100):
                # Add as heading
                heading = doc.add_heading(line, level=1 if line.isupper() else 2)
                headings_count += 1
                logger.debug(f"Added heading: {line[:30]}..." if len(line) > 30 else f"Added heading: {line}")
                # Continue to next line
                continue
            
            # Check if this is a numbered item (e.g., "1.", "a.", "i.")
            numbered_item_match = re.match(r'^(\d+|[a-z]|[ivxl]+)[\.\)]', line)
            
            if numbered_item_match:
                # This is a numbered item
                in_numbered_list = True
                list_items_count += 1
                
                # Check the type of numbering to determine the level
                if re.match(r'^\d+[\.\)]', line):
                    list_level = 0
                elif re.match(r'^[a-z][\.\)]', line):
                    list_level = 1
                else:
                    list_level = 2
                
                # Extract the content after the number
                content = line[line.find(' ')+1:] if ' ' in line else line
                
                # Add as a list item
                p = doc.add_paragraph(content, style='List Number')
                # Use Twips or integer for indentation (720 twips = 1/2 inch)
                p.paragraph_format.left_indent = Twips(720 * list_level)
                logger.debug(f"Added list item (level {list_level}): {content[:30]}..." if len(content) > 30 else f"Added list item (level {list_level}): {content}")
            else:
                # Regular paragraph
                in_numbered_list = False
                doc.add_paragraph(line)
                paragraphs_count += 1
                
                # Log significant paragraphs (like "WHEREAS" clauses)
                if line.startswith("WHEREAS") or line.startswith("NOW") or line.startswith("IN WITNESS"):
                    logger.debug(f"Added significant paragraph: {line[:30]}..." if len(line) > 30 else f"Added significant paragraph: {line}")
        
        # Add signature lines with proper spacing if the document ends with signatures
        if any("signature" in line.lower() for line in lines[-10:]):
            logger.debug("Detected signature section, adding proper spacing")
            for _ in range(3):
                doc.add_paragraph("")
        
        # Log document structure summary
        logger.info(f"Document structure: {headings_count} headings, {paragraphs_count} paragraphs, {list_items_count} list items")
        
        # Save to BytesIO object
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        conversion_time = time.time() - start_time
        logger.info(f"DOCX conversion completed in {conversion_time:.2f} seconds")
        
        return docx_file.getvalue()
        
    except Exception as e:
        logger.error(f"Error during DOCX conversion: {str(e)}", exc_info=True)
        # Create a simple error document
        error_doc = Document()
        error_doc.add_heading("Conversion Error", 0)
        error_doc.add_paragraph(f"An error occurred while converting the document: {str(e)}")
        error_doc.add_paragraph("Please try again or contact support.")
        
        error_file = BytesIO()
        error_doc.save(error_file)
        error_file.seek(0)
        
        return error_file.getvalue()